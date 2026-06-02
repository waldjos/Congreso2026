from docx import Document
import os
path = os.path.abspath(os.path.join('public', 'XXXVI CONGRESO NACIONAL DE UROLOGIA-1.docx'))
doc = Document(path)
print('Paragraphs:', len(doc.paragraphs))
for i, p in enumerate(doc.paragraphs[:120]):
    text = p.text.strip()
    if text:
        print(i, repr(text))
print('--- Tables:', len(doc.tables))
for ti, table in enumerate(doc.tables):
    print('Table', ti, 'rows', len(table.rows), 'cols', len(table.columns))
    for ri, row in enumerate(table.rows[:5]):
        print(' ', [cell.text.strip() for cell in row.cells])
