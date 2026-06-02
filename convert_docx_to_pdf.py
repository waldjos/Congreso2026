from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from xml.sax.saxutils import escape
import os

DOCX_FILENAME = os.path.join('public', 'XXXVI CONGRESO NACIONAL DE UROLOGIA-1.docx')
PDF_FILENAME = os.path.join('public', 'Programa-Congreso-2026.pdf')

def text_for_paragraph(paragraph):
    text = paragraph.text.strip()
    return escape(text) if text else ''


def normalize_cell_text(cell):
    return escape(' '.join(line.strip() for line in cell.text.splitlines() if line.strip()))


def main():
    if not os.path.exists(DOCX_FILENAME):
        raise FileNotFoundError(f"DOCX file not found: {DOCX_FILENAME}")

    doc = Document(DOCX_FILENAME)
    styles = getSampleStyleSheet()
    normal_style = styles['BodyText']
    heading_style = ParagraphStyle(
        'Heading', parent=styles['Heading1'], fontSize=16, leading=20, spaceAfter=10, spaceBefore=12
    )
    subheading_style = ParagraphStyle(
        'Subheading', parent=styles['Heading2'], fontSize=14, leading=18, spaceAfter=8, spaceBefore=10
    )

    flowables = [
        Paragraph('Programa - XXXVI Congreso Venezolano de Urología 2026', heading_style),
        Spacer(1, 0.2 * inch),
    ]

    for paragraph in doc.paragraphs:
        text = text_for_paragraph(paragraph)
        if not text:
            continue

        style_name = paragraph.style.name.lower() if paragraph.style else ''
        if 'heading 1' in style_name or 'title' in style_name:
            style = heading_style
        elif 'heading 2' in style_name or 'subtitle' in style_name:
            style = subheading_style
        else:
            style = normal_style

        flowables.append(Paragraph(text, style))
        flowables.append(Spacer(1, 0.1 * inch))

    for table_index, table in enumerate(doc.tables):
        if table_index == 0:
            flowables.append(Spacer(1, 0.2 * inch))
            flowables.append(Paragraph('Tabla de contenido', subheading_style))

        data = []
        for row in table.rows:
            data.append([normalize_cell_text(cell) for cell in row.cells])

        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e8eaf6')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.gray),
            ('BOX', (0, 0), (-1, -1), 0.25, colors.gray),
        ])
        reportlab_table = Table(data, colWidths=[4 * inch] * len(table.columns), repeatRows=1)
        reportlab_table.setStyle(table_style)
        flowables.append(reportlab_table)
        flowables.append(Spacer(1, 0.2 * inch))

    doc_pdf = SimpleDocTemplate(PDF_FILENAME, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    doc_pdf.build(flowables)
    print(f"Generated PDF: {PDF_FILENAME}")


if __name__ == '__main__':
    main()
